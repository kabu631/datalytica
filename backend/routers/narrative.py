import os
import uuid
import json
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List, Optional
from database import get_db
from models import Dataset, Report
from services.eda_service import compute_eda_full
from services.ai_service import AIService

router = APIRouter()

class GenerateNarrativeReq(BaseModel):
    dataset_id: int
    report_title: str
    tone: str
    focus_columns: Optional[List[str]] = []

@router.post("/generate")
def generate_narrative(req: GenerateNarrativeReq, db: Session = Depends(get_db)):
    """Generate an AI narrative report based on full EDA stats."""
    ds = db.query(Dataset).filter(Dataset.id == req.dataset_id).first()
    if not ds:
        raise HTTPException(status_code=404, detail="Dataset not found")
        
    eda_stats = compute_eda_full(req.dataset_id, db)["raw_stats"]
    
    if req.focus_columns:
        eda_stats["numeric_stats"] = {k: v for k, v in eda_stats["numeric_stats"].items() if k in req.focus_columns}
        eda_stats["categorical_stats"] = {k: v for k, v in eda_stats["categorical_stats"].items() if k in req.focus_columns}
        eda_stats.pop("correlations", None)
        eda_stats.pop("distributions", None)
        
    system_prompt = f"""You are a business report writer. Write in {req.tone} tone.
Structure:
  1. Executive Summary (2-3 sentences)
  2. Key Findings (5 bullet points, each citing a specific number)
  3. Notable Patterns (2-3 paragraphs)
  4. Recommendations (3 actionable items based only on the data)
 
STRICT RULES:
- Every number you write must appear in the stats JSON below
- Never write a sentence without a data reference
- Forbidden words: 'likely', 'probably', 'may', 'might', 'could suggest'
- If you cannot support a claim with the data, omit it
"""
    user_prompt = f"Dataset: {ds.name}\nStats:\n{json.dumps(eda_stats, default=str)}"
    
    with AIService(db) as ai:
        content = ai.complete(system_prompt=system_prompt, user_prompt=user_prompt)
        
    report = Report(
        dataset_id=req.dataset_id,
        title=req.report_title,
        content_markdown=content
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    
    return {
        "report_id": report.id,
        "title": report.title,
        "content_markdown": report.content_markdown
    }

class ExportRequest(BaseModel):
    report_id: int
    format: str

@router.post("/export")
def export_narrative(req: ExportRequest, db: Session = Depends(get_db)):
    """Export the markdown report into PDF, DOCX, or pure MD formats."""
    report = db.query(Report).filter(Report.id == req.report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
        
    os.makedirs("exports", exist_ok=True)
    filename = f"report_{uuid.uuid4().hex[:6]}.{req.format}"
    filepath = os.path.join("exports", filename)
    filepath = os.path.abspath(filepath)
    
    if req.format == "markdown":
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(report.content_markdown)
    elif req.format == "docx":
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(report.title, 0)
            for line in report.content_markdown.split('\n'):
                line = line.strip()
                if not line: continue
                
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('-') or line.startswith('*'):
                    doc.add_paragraph(line[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line)
            doc.save(filepath)
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx is not installed on the server.")
    elif req.format == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            flowables = [Paragraph(report.title, styles['Title']), Spacer(1, 12)]
            
            for line in report.content_markdown.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('### '):
                    flowables.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith('## '):
                    flowables.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('# '):
                    flowables.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('-') or line.startswith('*'):
                    flowables.append(Paragraph(line, styles['Bullet']))
                else:
                    line = line.replace('**', '').replace('*', '')
                    flowables.append(Paragraph(line, styles['Normal']))
                flowables.append(Spacer(1, 6))
            doc.build(flowables)
        except ImportError:
            raise HTTPException(status_code=500, detail="reportlab is not installed on the server.")
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use markdown, docx, or pdf")
        
    return FileResponse(path=filepath, filename=f"{report.title}.{req.format}")
